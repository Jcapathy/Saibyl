# PUBLIC INTERFACE
# ─────────────────────────────────────────────────────────
# process_document(document_id: UUID) -> DocumentContent
# get_project_text(project_id: UUID) -> str
# ─────────────────────────────────────────────────────────
from __future__ import annotations

import re
import tempfile
from pathlib import Path
from uuid import UUID

import fitz  # PyMuPDF
import structlog
from charset_normalizer import from_bytes
from pydantic import BaseModel

from app.core.database import get_supabase_admin

logger = structlog.get_logger()

MAX_DOC_SIZE = 50 * 1024 * 1024  # 50MB
MAX_PROJECT_SIZE = 200 * 1024 * 1024  # 200MB
MAX_LLM_CHARS = 50_000


class DocumentContent(BaseModel):
    document_id: UUID
    raw_text: str
    chunks: list[str]
    encoding: str
    page_count: int | None = None


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """Chunk text on sentence boundaries with overlap."""
    sentences = re.split(r"(?<=[.!?])\s+", text)
    chunks: list[str] = []
    current = ""

    for sentence in sentences:
        if len(current) + len(sentence) > chunk_size and current:
            chunks.append(current.strip())
            # Keep overlap from end of previous chunk
            overlap_text = current[-overlap:] if len(current) > overlap else current
            current = overlap_text + " " + sentence
        else:
            current = (current + " " + sentence).strip()

    if current.strip():
        chunks.append(current.strip())

    return chunks


def _extract_pdf(file_path: str) -> tuple[str, int]:
    doc = fitz.open(file_path)
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n\n".join(pages), len(pages)


def _extract_text(file_bytes: bytes, file_type: str) -> tuple[str, str, int | None]:
    """Extract text, detect encoding. Returns (text, encoding, page_count)."""
    if file_type == "pdf":
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp.flush()
            text, pages = _extract_pdf(tmp.name)
            Path(tmp.name).unlink(missing_ok=True)
        return text, "utf-8", pages

    # DOCX is a binary ZIP format — must be handled before charset detection
    if file_type == "docx":
        return _extract_docx(file_bytes), "utf-8", None

    # For text-based formats, detect encoding
    result = from_bytes(file_bytes)
    best = result.best()
    encoding = best.encoding if best else "utf-8"
    text = file_bytes.decode(encoding, errors="replace")
    return text, encoding, None


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file.

    Chain: python-docx → zipfile+XML (stdlib) → unstructured → error.
    The zipfile fallback is guaranteed to work on any valid DOCX since
    DOCX is just a ZIP archive containing word/document.xml."""
    import io as _io

    logger.info("docx_extraction_start", byte_length=len(file_bytes),
                header=file_bytes[:4].hex() if len(file_bytes) >= 4 else "short")

    # Method 1: python-docx
    try:
        import docx as _docx

        doc = _docx.Document(_io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if paragraphs:
            logger.info("docx_python_docx_ok", paragraphs=len(paragraphs))
            return "\n\n".join(paragraphs)
        # Try tables if no paragraphs
        rows = []
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
        if rows:
            logger.info("docx_python_docx_tables_ok", rows=len(rows))
            return "\n".join(rows)
        logger.warning("docx_python_docx_empty")
    except Exception as e:
        logger.warning("docx_python_docx_failed", error=str(e), error_type=type(e).__name__)

    # Method 2: stdlib zipfile + XML text extraction (guaranteed for valid DOCX)
    try:
        import zipfile
        import xml.etree.ElementTree as ET

        with zipfile.ZipFile(_io.BytesIO(file_bytes)) as zf:
            if "word/document.xml" not in zf.namelist():
                logger.warning("docx_zip_no_document_xml", files=zf.namelist()[:10])
            else:
                xml_bytes = zf.read("word/document.xml")
                root = ET.fromstring(xml_bytes)
                # DOCX namespace for text nodes
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                texts = []
                for paragraph in root.iter(f"{{{ns['w']}}}p"):
                    runs = paragraph.findall(f".//{{{ns['w']}}}t")
                    para_text = "".join(r.text or "" for r in runs).strip()
                    if para_text:
                        texts.append(para_text)
                if texts:
                    logger.info("docx_zipxml_ok", paragraphs=len(texts))
                    return "\n\n".join(texts)
                logger.warning("docx_zipxml_empty")
    except Exception as e:
        logger.warning("docx_zipxml_failed", error=str(e), error_type=type(e).__name__)

    # Method 3: unstructured (may not be installed)
    try:
        from unstructured.partition.docx import partition_docx

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp.flush()
            elements = partition_docx(filename=tmp.name)
            text = "\n\n".join(str(el) for el in elements)
            Path(tmp.name).unlink(missing_ok=True)
        if text.strip():
            return text
    except Exception as e:
        logger.warning("docx_unstructured_failed", error=str(e))

    logger.error("docx_all_methods_failed", byte_length=len(file_bytes))
    return "[Unable to extract text from this DOCX file]"


async def process_document(document_id: UUID) -> DocumentContent:
    """Download, extract, and chunk a document."""
    admin = get_supabase_admin()

    # Fetch document record
    doc_result = (
        admin.table("documents")
        .select("*")
        .eq("id", str(document_id))
        .single()
        .execute()
    )
    doc = doc_result.data

    # Update status to processing
    admin.table("documents").update(
        {"processing_status": "processing"}
    ).eq("id", str(document_id)).execute()

    try:
        # Download from Supabase Storage
        file_bytes = admin.storage.from_("project-media").download(doc["storage_path"])

        if len(file_bytes) > MAX_DOC_SIZE:
            raise ValueError(f"Document exceeds {MAX_DOC_SIZE // (1024*1024)}MB limit")

        text, encoding, page_count = _extract_text(file_bytes, doc["file_type"])
        chunks = chunk_text(text)

        # Update status to complete
        admin.table("documents").update(
            {"processing_status": "complete"}
        ).eq("id", str(document_id)).execute()

        logger.info(
            "document_processed",
            document_id=str(document_id),
            chars=len(text),
            chunks=len(chunks),
        )

        return DocumentContent(
            document_id=document_id,
            raw_text=text,
            chunks=chunks,
            encoding=encoding,
            page_count=page_count,
        )

    except Exception as e:
        admin.table("documents").update(
            {"processing_status": "failed", "error_message": str(e)}
        ).eq("id", str(document_id)).execute()
        logger.error("document_processing_failed", document_id=str(document_id), error=str(e))
        raise


async def get_project_text(project_id: UUID) -> str:
    """Get combined text of all processed documents in a project."""
    admin = get_supabase_admin()
    docs = (
        admin.table("documents")
        .select("id, storage_path, file_type")
        .eq("project_id", str(project_id))
        .eq("processing_status", "complete")
        .execute()
    )

    all_text: list[str] = []
    total_size = 0
    for doc in docs.data:
        file_bytes = admin.storage.from_("project-media").download(doc["storage_path"])
        total_size += len(file_bytes)
        if total_size > MAX_PROJECT_SIZE:
            logger.warning("project_text_size_limit", project_id=str(project_id))
            break
        text, _, _ = _extract_text(file_bytes, doc["file_type"])
        all_text.append(text)

    return "\n\n---\n\n".join(all_text)
